"""
Script pour générer le document Word de l'architecture du projet d'examens
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def set_cell_shading(cell, color):
    """Applique une couleur de fond à une cellule"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def add_styled_table(doc, headers, rows, header_color="1e3a5f"):
    """Crée un tableau stylisé"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        set_cell_shading(cell, header_color)

    # Données
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if r % 2 == 0:
                set_cell_shading(cell, "f2f2f2")

    return table

# ============================================================
# PAGE DE GARDE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ARCHITECTURE GLOBALE")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(30, 58, 95)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Plateforme de Gestion d'Examens")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("Frontend  •  Backend  •  IA")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(139, 69, 19)

for _ in range(4):
    doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("Février 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ============================================================
# SOMMAIRE
# ============================================================
doc.add_heading("Sommaire", level=1)
sommaire_items = [
    "1. Vue d'ensemble",
    "2. Frontend — React + Vite + TypeScript",
    "3. Backend — Node.js + Express",
    "4. Base de données — PostgreSQL",
    "5. IA — FastAPI + Python",
    "6. Flux de données principal",
    "7. Résumé des technologies",
]
for item in sommaire_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 1. VUE D'ENSEMBLE
# ============================================================
doc.add_heading("1. Vue d'ensemble", level=1)

doc.add_paragraph(
    "Ce projet est une plateforme complète de gestion d'examens composée de 3 grandes briques :"
)

overview_items = [
    ("Frontend (React)", "Interface utilisateur pour les professeurs, étudiants et administrateurs"),
    ("Backend (Node.js + Express)", "API REST servant de couche intermédiaire entre le frontend, la base de données et le module IA"),
    ("IA (FastAPI + Python)", "Moteur de correction automatique intelligent supportant QCM, calculs, questions courtes et longues"),
]

for title_text, desc_text in overview_items:
    p = doc.add_paragraph()
    run = p.add_run(f"• {title_text} : ")
    run.bold = True
    p.add_run(desc_text)

doc.add_paragraph()
doc.add_paragraph(
    "Le système gère 3 rôles utilisateurs : admin, professeur et etudiant. "
    "L'authentification est assurée par JWT côté backend Node.js."
)

doc.add_paragraph()

# Schéma textuel de l'architecture
doc.add_heading("Schéma d'architecture", level=2)

schema = doc.add_paragraph()
schema.paragraph_format.space_before = Pt(6)
schema_text = """┌─────────────────────┐       ┌──────────────────────────┐       ┌─────────────────────┐
│   FRONTEND          │       │   BACKEND                │       │   IA                │
│   React + Vite      │──────▶│   Node.js + Express      │──────▶│   FastAPI + Python   │
│   TypeScript        │ REST  │   :3000                  │ HTTP  │   :8000             │
│   shadcn/ui         │◀──────│   JWT + bcrypt           │◀──────│   Correction auto   │
│   Tailwind CSS      │       │   Prisma/Knex            │       │   Ollama LLM        │
└─────────────────────┘       └────────────┬─────────────┘       └─────────────────────┘
                                           │
                                           │ SQL
                                           ▼
                              ┌──────────────────────────┐
                              │   PostgreSQL              │
                              │   Base de données         │
                              │   (9 tables principales)  │
                              └──────────────────────────┘"""
run = schema.add_run(schema_text)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# ============================================================
# 2. FRONTEND
# ============================================================
doc.add_heading("2. Frontend — React + Vite + TypeScript", level=1)

doc.add_heading("2.1. Stack technique", level=2)
add_styled_table(doc,
    ["Aspect", "Détail"],
    [
        ["Framework", "React 18 + TypeScript"],
        ["Build tool", "Vite"],
        ["UI Components", "shadcn/ui (basé sur Radix UI)"],
        ["CSS", "Tailwind CSS"],
        ["State / Data fetching", "TanStack React Query"],
        ["Routing", "React Router DOM v6"],
        ["Auth côté client", "JWT stocké en localStorage, envoyé via Bearer Token"],
    ]
)

doc.add_paragraph()
doc.add_heading("2.2. Pages par rôle", level=2)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Professeur :")
run.bold = True
p.add_run(" Dashboard, CreateExam, Questions, Exams, ExamSubmissions, Corrections, Classes, Analytics")

p = doc.add_paragraph()
run = p.add_run("Etudiant :")
run.bold = True
p.add_run(" MyExams, TakeExam, MyResults")

p = doc.add_paragraph()
run = p.add_run("Admin :")
run.bold = True
p.add_run(" Users, Subjects, Levels, SendNotifications")

p = doc.add_paragraph()
run = p.add_run("Commun :")
run.bold = True
p.add_run(" Auth, Profile, Settings, Notifications")

doc.add_paragraph()
doc.add_heading("2.3. Structure des dossiers", level=2)

structure = """src/
├── components/          # Composants réutilisables
│   ├── layout/          # Layout (Sidebar, Header)
│   └── ui/              # Composants shadcn/ui
├── hooks/               # Hooks custom (useAuth, useTheme, use-toast)
├── lib/                 # Utilitaires (appels API, correction auto)
├── pages/               # Pages de l'application
│   └── admin/           # Pages admin
└── services/            # (à créer) Appels HTTP vers le backend Node.js"""

p = doc.add_paragraph()
run = p.add_run(structure)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_heading("2.4. Communication avec le backend", level=2)
doc.add_paragraph(
    "Le frontend communique avec le backend Node.js via des requêtes HTTP REST "
    "(fetch ou Axios). Le token JWT est inclus dans le header Authorization: Bearer <token>. "
    "Le frontend n'a plus de connexion directe à la base de données (contrairement à l'ancien modèle Supabase)."
)

doc.add_page_break()

# ============================================================
# 3. BACKEND
# ============================================================
doc.add_heading("3. Backend — Node.js + Express", level=1)

doc.add_heading("3.1. Stack technique", level=2)
add_styled_table(doc,
    ["Aspect", "Détail"],
    [
        ["Runtime", "Node.js"],
        ["Framework HTTP", "Express.js"],
        ["Port", ":3000"],
        ["Authentification", "JWT (jsonwebtoken) + bcrypt pour le hachage des mots de passe"],
        ["ORM / Query Builder", "Prisma ou Knex.js"],
        ["Validation", "Joi ou Zod"],
        ["CORS", "cors middleware"],
    ],
    header_color="8b4513"
)

doc.add_paragraph()
doc.add_heading("3.2. Architecture en couches", level=2)

layers = [
    ("Routes / Router", "Définition des endpoints REST (GET, POST, PUT, DELETE)"),
    ("Middlewares", "Authentification JWT, vérification des rôles (admin/professeur/etudiant), validation des données d'entrée"),
    ("Controllers", "Réception des requêtes, appel aux services, formatage des réponses"),
    ("Services", "Logique métier : création d'examens, soumission de copies, calcul de notes, appel à l'API IA"),
    ("ORM (Prisma/Knex)", "Requêtes SQL vers PostgreSQL, migrations de schéma"),
]

for title_text, desc_text in layers:
    p = doc.add_paragraph()
    run = p.add_run(f"• {title_text} : ")
    run.bold = True
    p.add_run(desc_text)

doc.add_paragraph()
doc.add_heading("3.3. Structure des dossiers (proposée)", level=2)

backend_structure = """backend/
├── src/
│   ├── config/              # Configuration (DB, JWT, env)
│   ├── middlewares/          # auth.js, roles.js, validate.js
│   ├── routes/               # Définition des routes
│   │   ├── auth.routes.js
│   │   ├── exams.routes.js
│   │   ├── questions.routes.js
│   │   ├── submissions.routes.js
│   │   ├── classes.routes.js
│   │   ├── users.routes.js
│   │   └── notifications.routes.js
│   ├── controllers/          # Logique de contrôle
│   ├── services/             # Logique métier
│   ├── models/               # Modèles Prisma ou schémas Knex
│   └── utils/                # Helpers (hash, token, etc.)
├── prisma/
│   └── schema.prisma         # Schéma de la BDD
├── package.json
└── .env                      # Variables d'environnement"""

p = doc.add_paragraph()
run = p.add_run(backend_structure)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_heading("3.4. Endpoints API principaux", level=2)

add_styled_table(doc,
    ["Méthode", "Endpoint", "Description"],
    [
        ["POST", "/api/auth/register", "Inscription d'un utilisateur"],
        ["POST", "/api/auth/login", "Connexion (retourne un JWT)"],
        ["GET", "/api/users/me", "Profil de l'utilisateur connecté"],
        ["GET", "/api/exams", "Liste des examens"],
        ["POST", "/api/exams", "Créer un examen (professeur)"],
        ["GET", "/api/exams/:id", "Détail d'un examen"],
        ["POST", "/api/exams/:id/submit", "Soumettre une copie (étudiant)"],
        ["GET", "/api/exams/:id/submissions", "Soumissions d'un examen (professeur)"],
        ["GET", "/api/questions", "Banque de questions"],
        ["POST", "/api/questions", "Créer une question"],
        ["GET", "/api/classes", "Liste des classes"],
        ["POST", "/api/classes", "Créer une classe"],
        ["GET", "/api/notifications", "Notifications de l'utilisateur"],
        ["GET", "/api/analytics", "Statistiques (professeur)"],
        ["POST", "/api/corrections/:submissionId", "Déclencher la correction IA"],
    ],
    header_color="8b4513"
)

doc.add_paragraph()
doc.add_heading("3.5. Appel à l'API IA", level=2)
doc.add_paragraph(
    "Lorsqu'un étudiant soumet une copie, le backend Node.js :"
)
steps = [
    "Sauvegarde les réponses dans PostgreSQL (tables submissions + answers)",
    "Prépare les données au format attendu par l'API IA (evaluation + copie)",
    "Envoie une requête HTTP POST vers http://localhost:8000/api/corriger-copie",
    "Reçoit le résultat (notes + feedbacks par question)",
    "Met à jour PostgreSQL avec les scores et feedbacks",
    "Retourne le résultat au frontend",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_page_break()

# ============================================================
# 4. BASE DE DONNÉES
# ============================================================
doc.add_heading("4. Base de données — PostgreSQL", level=1)

doc.add_heading("4.1. Tables principales", level=2)
add_styled_table(doc,
    ["Table", "Description", "Colonnes clés"],
    [
        ["profiles", "Utilisateurs de la plateforme", "id, email, full_name, avatar_url"],
        ["user_roles", "Rôles des utilisateurs", "user_id, role (admin | professeur | etudiant)"],
        ["subjects", "Matières", "id, name, description, color"],
        ["levels", "Niveaux d'enseignement", "id, name, description"],
        ["classes", "Classes", "id, name, level_id, created_by"],
        ["class_students", "Inscriptions élèves dans les classes", "class_id, student_id"],
        ["questions", "Banque de questions", "id, question_text, question_type, options, correct_answer, points, difficulty, subject_id"],
        ["exams", "Examens", "id, title, description, status, duration_minutes, total_points, class_id, subject_id, level_id"],
        ["exam_questions", "Questions rattachées à un examen", "exam_id, question_id, order_index, points"],
        ["submissions", "Soumissions des étudiants", "id, exam_id, student_id, status, score, submitted_at"],
        ["answers", "Réponses individuelles", "id, submission_id, question_id, answer_text, is_correct, points_awarded, feedback"],
        ["notifications", "Notifications", "id, user_id, title, message, type, is_read"],
    ],
    header_color="2d5016"
)

doc.add_paragraph()
doc.add_heading("4.2. Enums", level=2)
add_styled_table(doc,
    ["Enum", "Valeurs possibles"],
    [
        ["app_role", "admin, professeur, etudiant"],
        ["exam_status", "brouillon, publie, corrige, archive"],
        ["question_type", "qcm, vrai_faux, reponse_courte, redaction"],
        ["difficulty_level", "facile, moyen, difficile"],
    ],
    header_color="2d5016"
)

doc.add_paragraph()
doc.add_heading("4.3. Relations principales", level=2)
relations = [
    "exams → classes, subjects, levels (FK)",
    "exam_questions → exams, questions (FK)",
    "questions → subjects (FK)",
    "submissions → exams (FK)",
    "answers → submissions, questions (FK)",
    "classes → levels (FK)",
    "class_students → classes (FK)",
]
for rel in relations:
    doc.add_paragraph(f"• {rel}")

doc.add_page_break()

# ============================================================
# 5. IA
# ============================================================
doc.add_heading("5. IA — FastAPI + Python", level=1)

doc.add_heading("5.1. Stack technique", level=2)
add_styled_table(doc,
    ["Aspect", "Détail"],
    [
        ["Framework API", "FastAPI (Python)"],
        ["Port", ":8000"],
        ["Architecture", "Strategy Pattern (1 correcteur par type de question)"],
        ["Documentation", "Swagger auto-générée sur /docs"],
    ],
    header_color="5c1a5c"
)

doc.add_paragraph()
doc.add_heading("5.2. Moteur de correction (CorrectionEngine)", level=2)
doc.add_paragraph(
    "Le moteur principal route chaque question vers le correcteur approprié selon son type :"
)

add_styled_table(doc,
    ["Type", "Correcteur", "Technologie", "Méthode"],
    [
        ["QCM", "QCMCorrector", "Comparaison directe", "Match exact de la réponse attendue"],
        ["Calcul", "CalculCorrector", "SymPy", "Vérification symbolique/numérique avec tolérance configurable"],
        ["Courte", "CourteCorrector", "Sentence Transformers (paraphrase-multilingual-mpnet-base-v2)", "Similarité sémantique par embeddings"],
        ["Longue", "LongueCorrector", "Ollama (llama3.2 / mistral)", "Évaluation multicritères par prompt LLM"],
    ],
    header_color="5c1a5c"
)

doc.add_paragraph()
doc.add_heading("5.3. Endpoints API IA", level=2)
add_styled_table(doc,
    ["Méthode", "Endpoint", "Description"],
    [
        ["POST", "/api/corriger-question", "Corrige une question unique"],
        ["POST", "/api/corriger-copie", "Corrige une copie complète d'un élève"],
        ["POST", "/api/corriger-copies", "Corrige plusieurs copies (batch)"],
        ["POST", "/api/extraire-questions-pdf", "Extraction automatique de questions depuis un PDF"],
        ["GET", "/health", "Statut de santé des modules"],
        ["GET", "/api/types-questions", "Liste des types de questions supportés"],
        ["GET", "/api/matieres", "Liste des matières supportées"],
    ],
    header_color="5c1a5c"
)

doc.add_paragraph()
doc.add_heading("5.4. Structure des dossiers", level=2)
ia_structure = """ia/
├── ia_correction/
│   ├── __init__.py
│   ├── api.py               # Serveur FastAPI et routes
│   ├── config.py             # Configuration (seuils, modèles, timeouts)
│   ├── engine.py             # CorrectionEngine (routeur principal)
│   ├── models.py             # Modèles Pydantic (Question, Réponse, Résultat)
│   └── correcteurs/
│       ├── base_corrector.py  # Classe abstraite BaseCorrector
│       ├── qcm_corrector.py   # Correction QCM
│       ├── calcul_corrector.py # Correction calculs (SymPy)
│       ├── courte_corrector.py # Correction courte (Sentence Transformers)
│       └── longue_corrector.py # Correction longue (Ollama LLM)
├── tests/
│   └── test_correction.py
├── requirements.txt
└── README.md"""

p = doc.add_paragraph()
run = p.add_run(ia_structure)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_heading("5.5. Configuration IA", level=2)
add_styled_table(doc,
    ["Paramètre", "Valeur", "Description"],
    [
        ["SENTENCE_TRANSFORMER_MODEL", "paraphrase-multilingual-mpnet-base-v2", "Modèle pour les embeddings sémantiques"],
        ["OLLAMA_MODEL", "llama3.2", "Modèle LLM pour les questions longues"],
        ["OLLAMA_HOST", "http://localhost:11434", "Adresse du serveur Ollama"],
        ["SIMILARITE_EXCELLENTE", "0.85", "Seuil pour ~95% des points"],
        ["SIMILARITE_BONNE", "0.70", "Seuil pour ~70% des points"],
        ["TOLERANCE_CALCUL", "0.01", "Tolérance numérique pour les calculs"],
        ["TIMEOUT_LLM", "60s", "Timeout pour les appels LLM"],
    ],
    header_color="5c1a5c"
)

doc.add_page_break()

# ============================================================
# 6. FLUX DE DONNÉES
# ============================================================
doc.add_heading("6. Flux de données principal", level=1)

doc.add_heading("6.1. Soumission et correction d'un examen", level=2)

flux = """┌──────────────┐     ┌──────────────────────┐     ┌───────────────────┐
│  ÉTUDIANT    │     │  BACKEND Node.js     │     │  API IA FastAPI   │
│  (Frontend)  │     │  Express :3000       │     │  :8000            │
└──────┬───────┘     └──────────┬───────────┘     └─────────┬─────────┘
       │                       │                            │
       │  1. POST /submit      │                            │
       │  (réponses + JWT)     │                            │
       │──────────────────────▶│                            │
       │                       │  2. Vérif JWT + rôle       │
       │                       │  3. Sauvegarde PostgreSQL  │
       │                       │     (submissions + answers)│
       │                       │                            │
       │                       │  4. POST /corriger-copie   │
       │                       │──────────────────────────▶│
       │                       │                            │
       │                       │                 5. Correction par type :
       │                       │                    QCM → comparaison
       │                       │                    Calcul → SymPy
       │                       │                    Courte → embeddings
       │                       │                    Longue → Ollama LLM
       │                       │                            │
       │                       │  6. Résultat (notes +      │
       │                       │◀─────────────feedbacks)────│
       │                       │                            │
       │                       │  7. MAJ PostgreSQL         │
       │                       │     (score, feedback,      │
       │  8. Réponse JSON      │     statut = "corrige")    │
       │◀──────────────────────│                            │
       │                       │                            │
       ▼                       ▼                            ▼
  Affichage résultats"""

p = doc.add_paragraph()
run = p.add_run(flux)
run.font.name = 'Consolas'
run.font.size = Pt(7.5)

doc.add_page_break()

# ============================================================
# 7. RÉSUMÉ
# ============================================================
doc.add_heading("7. Résumé des technologies", level=1)

add_styled_table(doc,
    ["Couche", "Technologies"],
    [
        ["Frontend", "React 18, TypeScript, Vite, Tailwind CSS, shadcn/ui (Radix UI), React Query, React Router DOM"],
        ["Backend", "Node.js, Express.js, JWT (jsonwebtoken), bcrypt, Prisma ou Knex.js"],
        ["Base de données", "PostgreSQL (12 tables, 4 enums)"],
        ["IA - Correction", "FastAPI (Python), SymPy, Sentence Transformers, Ollama (LLM local)"],
        ["IA - Extraction PDF", "PyPDF2, Ollama (llama3.2 / mistral)"],
        ["Infra locale", "Backend :3000, PostgreSQL :5432, FastAPI :8000, Ollama :11434"],
    ],
    header_color="333333"
)

doc.add_paragraph()

doc.add_heading("Ports de communication", level=2)
add_styled_table(doc,
    ["Service", "Port", "Protocole"],
    [
        ["Frontend (dev)", ":5173", "HTTP (Vite dev server)"],
        ["Backend Node.js", ":3000", "HTTP REST + JWT"],
        ["PostgreSQL", ":5432", "TCP/SQL"],
        ["API IA FastAPI", ":8000", "HTTP REST"],
        ["Ollama LLM", ":11434", "HTTP"],
    ],
    header_color="333333"
)

# ============================================================
# SAUVEGARDE
# ============================================================
output_path = r"d:\projet api sgstock\projet_iut\Architecture_Projet_Examens.docx"
doc.save(output_path)
print(f"Document Word généré : {output_path}")
